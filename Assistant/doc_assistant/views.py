from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
import spacy
from textblob import TextBlob
from langdetect import detect
from spacy.matcher import Matcher, PhraseMatcher
from django.conf import settings
import docx
import PyPDF2
from io import BytesIO

from doc_assistant.serializers import (
        DocumentSerializer,
        SuggestionSerializer,
        DocumentUploadSerializer,
        UserDocumentListSerializer
        )
from doc_assistant.models import (
        Document,
        Suggestion
        )


class DocumentProcessor:
    """
    Handles all document processing logic: extraction and improvement
    """
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm", disable=['ner'])
        if 'senter' not in self.nlp.pipe_names:
            self.nlp.add_pipe('sentencizer')

        # setup matchers
        self.matcher = Matcher(self.nlp.vocab)
        self.phrase_matcher = PhraseMatcher(self.nlp.vocab)

        passive_pattern = [
                {'DEP': 'auxpass'},
                {'TAG': 'VBN'}
                ]
        
        self.matcher.add('Passive', [passive_pattern])

    def extract_content(self, file):
        """
        Extract content from various file formats with enhanced error handling
        """
        if not file:
            raise ValueError("No file provided")

        file_extension = file.name.split('.')[-1].lower()

        try:
            if file_extension == 'txt':
                content = file.read().decode('utf-8', errors='replace')
            elif file_extension == 'docx':
                doc = docx.Document(file)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                content = '\n\n'.join(paragraphs)
            elif file_extension == 'pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text.strip())
                content = '\n\n'.join(text_parts)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            return self.validate_content(content)

        except Exception as e:
            raise ValueError(f"Error extracting content: {str(e)}")

    def validate_content(self, content):
        """
        Validate content length and language
        """
        if not content or len(content.strip()) == 0:
            raise ValueError("Extracted content is empty")

        try:
            if detect(content) != 'en':
                raise ValueError("Only English documents are supported")
        except Exception as e:
            raise ValueError(f"Language detection failed: {str(e)}")

        return content.strip()

    def improve_document(self, content):
        """
        Improve document content using NLP techniques
        """
        # initial spell check and basic grammar correction
        blob = TextBlob(content)
        corrected_text = str(blob.correct())

        # advanced processinf with SpaCy
        doc = self.nlp(corrected_text)
        improved_sentences = []

        for sent in doc.sents:
            improved_sent = self.improve_sentence(sent)
            improved_sentences.append(improved_sent)

        improved_text = ' '.join(improved_sentences)
        return improved_text

    def improve_sentence(self, sent):
        """
        Apply various improvements to a single sentence
        """
        text = sent.text.strip()

        # convert passive voice to active
        matches = self.matcher(sent)
        if matches:
            text = self.convert_to_active(sent)

        # fix capitalization
        if text and not text[0].isupper():
            text = text[0].upper() + text[1:]

        # ensure proper spacing and punctuation
        text = ' '.join(text.split())

        return text + ' '

    def convert_to_active(self, sent):
        """
        Convert passive voice to active voice
        """
        subject, agent, verb = None

        for token in sent:
            if token.dep_ == "nsubjpass":
                subject = token
            elif token.dep_ == "agent":
                agent = list(token.children)[0] if list(token.children) else None
            elif token.dep_ == "auxpass":
                verb = token.head

        if subject and verb:
            if agent:
                return f"{agent.text} {verb.text} {subject.text}"

        return sent.text


class DocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = DocumentSerializer
    document_processor = DocumentProcessor()

    def get_queryset(self):
        return Document.objects.filter(user=self.request.user)

    def get_serializer_class(self):
        if self.action == 'upload':
            return DocumentUploadSerializer
        return DocumentSerializer

    @action(detail=False, methods=['POST'])
    def upload(self, request):
        try:
            serializer = DocumentUploadSerializer(data=request.data)
            serializer.is_valid(raise_exception=True)

            file = serializer.validated_data['document']
            content = self.document_processor.extract_content(file)
            document = self.create_document(content, file.name)
            return Response(
                    DocumentSerializer(document).data,
                    status=status.HTTP_201_CREATED
                    )
        except Exception as e:
            print(f"Upload error: {str(e)}")
            return Response(
                    {'error': f'Upload failed: {str(e)}'},
                    status=status.HTTP_400_BAD_REQUEST
                    )
       
    def create_document(self, content, filename):
        """
        Create and process a new document
        """
        document = Document.objects.create(
                user=self.request.user,
                title=filename,
                original_content=content,
                file_type=filename.split('.')[-1].lower()
                )

        improved_content = self.document_processor.improve_document(content)
        document.improved_content = improved_content
        document.save()

        return document

    @action(detail=True, methods=['GET'])
    def compare(self, request, pk=None):
        document = self.get_object()
        return Response({
            'original_content': document.original_content,
            'improved_content': document.improved_content
            })

    @action(detail=True, methods=['POST'])
    def export(self, request, pk=None):
        document = self.get_object()

        try:
            # create new word document
            doc = docx.Document(settings.WORD_TEMPLATE_PATH)

            # add improved content
            doc.add_paragraph(document.improved_content)

            # add analysis
            doc.add_paragraph("\nDocument Analysis:", style='Heading 1')

            # sentiment analysis
            blob = TextBlob(document.improved_content)
            sentiment = blob.sentiment
            sentiment_text = 'Positive' if sentiment.polarity > 0 else 'Negative' if sentiment.polarity < 0 else 'Neutral'
            doc.add_paragraph(f"Sentiment: {sentiment_text}")

            # structural analysis
            spacy_doc = self.document_processor.nlp(document.improved_content)
            sentences = list(spacy_doc.sents)
            doc.add_paragraph(f"Total Sentences: {len(sentences)}")
            doc.add_paragraph(f"Average Sentence Length: {len(spacy_doc) / len(sentences):.lf} words")

            # save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = Response(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
            response['Content-Disposition'] = f'attachment; filename="{document.title}_improved.docx"'
            return response

        except Exception as e:
            return Response(
                    {'error': f'Export failed: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
